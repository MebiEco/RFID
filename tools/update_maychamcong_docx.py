#!/usr/bin/env python3
"""Tao Maychamcong.docx — giao thuc day du cho Backend."""
from docx import Document
from docx.shared import Pt, RGBColor
import shutil
from pathlib import Path
from datetime import date

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Maychamcong.docx"


def para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def json_block(doc, text):
    para(doc, "json", bold=True, size=10)
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.size = Pt(11)


def num(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for r in p.runs:
        r.font.size = Pt(11)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val


def build_document():
    doc = Document()
    doc.add_heading("Giao thức truyền thông — Máy chấm công RFID", 0)
    doc.add_heading("Backend Integration Spec (bắt buộc đọc)", level=1)

    para(doc, f"Phiên bản: {date.today().isoformat()}")
    para(doc, "Nguồn firmware: ESP32-S3 RFID Scanner — app_azure.c, app_rfid.c, scan_log.c, card_profile.c")
    para(doc, "Mọi giao tiếp qua Azure IoT Hub. Backend KHÔNG kết nối MQTT trực tiếp bằng DeviceId của thiết bị.")

    # ── 0. Backend phải làm gì (tóm tắt) ──
    doc.add_heading("0. Backend phải làm gì — tóm tắt nhanh", level=1)
    add_table(doc, ["Việc", "Backend làm", "Thiết bị trả về"], [
        ("Nhận chấm công", "Subscribe telemetry từ IoT Hub", "Gửi Code 601/602 + Index + TimeStamp"),
        ("Đăng ký thẻ", "Gọi Direct Method Control Code 603", "Response 200 + telemetry 603 (Index) trên Hub"),
        ("Xóa thẻ", "Gọi Direct Method Control Code 604", "Response 200 + telemetry 604 (Index) trên Hub"),
        ("Đồng bộ thiếu message", "Gọi 605 kèm LastIdx* / Missing[]", "Response 605 + telemetry bù 601/602/..."),
        ("OTA", "Gọi 606 + Url", "Response 200, tải firmware"),
    ])
    bullet(doc, "Lưu DB theo (DeviceId, Code, Index) — dedupe bắt buộc khi sync/flush.")
    bullet(doc, "Theo dõi 3 max Index riêng: LastIdxSwipe (602), LastIdxUnkn (601), LastIdxAdmin (603/604).")

    # ── 1. Tổng quan ──
    doc.add_heading("1. Hai hướng giao tiếp", level=1)
    add_table(doc, ["Hướng", "Cơ chế Azure", "Ai khởi tạo"], [
        ("Thiết bị → Backend", "Telemetry — topic devices/{DeviceId}/messages/events/", "Thiết bị (quét thẻ, xác nhận admin)"),
        ("Backend → Thiết bị", "Direct Method — method name Control", "Backend (603/604/605/606)"),
    ])
    bullet(doc, "Backend KHÔNG gửi ACK telemetry xuống thiết bị.")
    bullet(doc, "605 là lệnh từ Backend → Thiết bị; thiết bị KHÔNG gửi Code 605 lên telemetry.")

    doc.add_heading("1.1. Thông số MQTT thiết bị (tham khảo)", level=2)
    add_table(doc, ["Tham số", "Giá trị"], [
        ("Broker", "mqtts://{IoTHubHostName}:8883"),
        ("Client ID", "{DeviceId}"),
        ("Username", "{IoTHubHostName}/{DeviceId}/?api-version=2021-04-12"),
        ("Password", "SAS Token (cấu hình trên thiết bị)"),
        ("Keepalive", "240 giây | QoS publish: 1"),
    ])

    # ── 2. Telemetry ──
    doc.add_heading("2. Thiết bị → Backend: Telemetry (Backend NHẬN)", level=1)
    para(doc, "Topic thiết bị publish:", bold=True)
    json_block(doc, "devices/{DeviceId}/messages/events/")
    para(doc, "Backend route message qua Event Hub / Service Bus / Azure Function tùy kiến trúc.")

    doc.add_heading("2.1. Cấu trúc JSON", level=2)
    json_block(doc, """{
    "Code": 602,
    "Index": 42,
    "TimeStamp": 1765036309,
    "Data": {
        "DeviceName": "RFID_Scanner",
        "UID": "09F8D21E",
        "Name": "Nguyen Van A",
        "ID": "NV01"
    }
}""")

    add_table(doc, ["Trường", "Backend xử lý"], [
        ("Code", "601=thẻ lạ | 602=chấm công | 603/604=xác nhận admin"),
        ("Index", "Số thứ tự — 3 bộ đếm riêng (mục 2.3). Dùng dedupe + phát hiện gap"),
        ("TimeStamp", "Unix UTC (giây). Code 602: dùng làm giờ chấm công"),
        ("Data.UID", "Hex không dấu ':'"),
        ("Data.Name, Data.ID", "Rỗng nếu Code 601"),
        ("Data.DeviceName", "Luôn RFID_Scanner"),
    ])

    doc.add_heading("2.2. Bảng Code telemetry", level=2)
    add_table(doc, ["Code", "Khi nào", "Backend action"], [
        ("601", "Quét thẻ chưa đăng ký", "Cảnh báo thẻ lạ; có thể gọi 603 đăng ký"),
        ("602", "Quét thẻ đã đăng ký", "Ghi chấm công: TimeStamp, UID, Name, ID"),
        ("603", "Sau lưu thẻ (portal/LCD hoặc DM 603 thành công)", "Xác nhận profile đã cập nhật; nhận telemetry 603 + Index"),
        ("604", "Sau xóa thẻ (portal/LCD hoặc DM 604 thành công)", "Xác nhận profile đã xóa; nhận telemetry 604 + Index"),
    ])

    doc.add_heading("2.3. Index — Backend PHẢI lưu 3 counter riêng", level=2)
    add_table(doc, ["Loại", "Code", "Cột DB gợi ý", "NVS trên thiết bị"], [
        ("Quét đã ĐK", "602", "last_idx_swipe", "idx_swipe"),
        ("Quét thẻ lạ", "601", "last_idx_unkn", "idx_unkn"),
        ("Admin lưu/xóa", "603, 604", "last_idx_admin", "idx_admin"),
    ])
    para(doc, "KHÔNG gộp Index 601 và 602 vào một dãy. Gap detection làm riêng từng loại.")

    doc.add_heading("2.4. Ví dụ telemetry", level=2)
    para(doc, "601 — thẻ lạ:", bold=True)
    json_block(doc, """{"Code":601,"Index":5,"TimeStamp":1765036309,
    "Data":{"DeviceName":"RFID_Scanner","UID":"A1B2C3D4","Name":"","ID":""}}""")
    para(doc, "602 — chấm công:", bold=True)
    json_block(doc, """{"Code":602,"Index":88,"TimeStamp":1765036310,
    "Data":{"DeviceName":"RFID_Scanner","UID":"09F8D21E","Name":"Nguyen","ID":"NV01"}}""")
    para(doc, "603 — xác nhận lưu thẻ (portal/LCD/DM):", bold=True)
    json_block(doc, """{"Code":603,"Index":3,"TimeStamp":1765036311,
    "Data":{"DeviceName":"RFID_Scanner","UID":"09F8D21E","Name":"Nguyen","ID":"NV01"}}""")
    para(doc, "604 — xác nhận xóa thẻ:", bold=True)
    json_block(doc, """{"Code":604,"Index":4,"TimeStamp":1765036312,
    "Data":{"DeviceName":"RFID_Scanner","UID":"09F8D21E","Name":"Nguyen","ID":"NV01"}}""")

    doc.add_heading("2.5. Telemetry 603/604 — khi nào Backend nhận", level=2)
    bullet(doc, "Quẹt thẻ: chỉ telemetry 601 hoặc 602 (không có 603/604).")
    bullet(doc, "Lưu/xóa thẻ trên portal web hoặc màn hình cảm ứng: thiết bị gửi telemetry 603 hoặc 604 kèm Index (idx_admin).")
    bullet(doc, "Backend gọi Direct Method 603/604 thành công: thiết bị CŨNG gửi telemetry 603/604 (cùng Index với dòng ghi rfid_log.csv).")
    bullet(doc, "Direct Method response 603/604 KHÔNG chứa trường Index — Backend lấy Index từ telemetry Hub hoặc đối chiếu nhật ký trên thiết bị.")

    # ── 3. Direct Method ──
    doc.add_heading("3. Backend → Thiết bị: Direct Method Control", level=1)
    para(doc, "Backend gọi: Azure REST/SDK invoke-device-method, method name = Control")
    add_table(doc, ["Thuộc tính", "Giá trị"], [
        ("Method name", "Control (khuyến nghị)"),
        ("Phân biệt lệnh", "Trường Code trong JSON body"),
        ("Timeout", "30 giây"),
        ("Topic thiết bị trả lời", "$iothub/methods/res/{httpStatus}/?$rid={requestId}"),
    ])

    doc.add_heading("3.1. Code 603 — Ghi / sửa thẻ", level=2)
    para(doc, "Backend GỬI:", bold=True)
    json_block(doc, """{
    "Code": 603,
    "TimeStamp": 1765036309,
    "Data": { "UID": "09F8D21E", "Name": "Nguyen", "ID": "NV01" }
}""")
    add_table(doc, ["Trường", "Bắt buộc"], [
        ("Data.UID", "Có"),
        ("Data.Name, Data.ID", "Không (thiếu → lưu rỗng)"),
    ])
    para(doc, "Thiết bị TRẢ (Direct Method response):", bold=True)
    json_block(doc, """{"status":200,"payload":{"Code":603,"TimeStamp":...,"Message":"Updated UID 09F8D21E"}}""")
    para(doc, "Thiết bị GỬI THÊM lên Hub (telemetry, cùng lúc hoặc ngay sau response):", bold=True)
    json_block(doc, """{"Code":603,"Index":3,"TimeStamp":...,"Data":{"DeviceName":"RFID_Scanner","UID":"09F8D21E","Name":"Nguyen","ID":"NV01"}}""")
    bullet(doc, "Index telemetry 603 = idx_admin trên thiết bị (cùng số ghi vào rfid_log.csv).")
    para(doc, "Lỗi: HTTP 500 Save failed. Thiếu UID → timeout 30s, không có response.")

    doc.add_heading("3.2. Code 604 — Xóa thẻ", level=2)
    para(doc, "Backend GỬI:", bold=True)
    json_block(doc, """{"Code":604,"TimeStamp":1765036309,"Data":{"UID":"09F8D21E"}}""")
    para(doc, "Thiết bị TRẢ (luôn 200):", bold=True)
    json_block(doc, """{"status":200,"payload":{"Code":604,"Message":"Deleted or absent UID 09F8D21E"}}""")
    para(doc, "Thiết bị GỬI THÊM lên Hub (telemetry):", bold=True)
    json_block(doc, """{"Code":604,"Index":4,"TimeStamp":...,"Data":{"DeviceName":"RFID_Scanner","UID":"09F8D21E","Name":"Nguyen","ID":"NV01"}}""")
    bullet(doc, "Name/ID trong telemetry 604 = giá trị profile TRƯỚC khi xóa (lookup trên SD).")

    doc.add_heading("3.3. Code 605 — Đồng bộ (QUAN TRỌNG)", level=2)
    para(doc, "Mục đích: Backend phát hiện thiếu Index → gọi 605 → thiết bị gửi lại telemetry bù.", bold=True)

    para(doc, "Thiết bị thực hiện 2 bước:", bold=True)
    num(doc, "Flush /sdcard/az_pend.bin — message chưa publish Hub.")
    num(doc, "Đọc /sdcard/rfid_log.csv — gửi lại Index backend thiếu (kể cả đã gửi Hub trước đó).")

    para(doc, "A) Backend GỬI — có đối chiếu (KHUYẾN NGHỊ):", bold=True)
    json_block(doc, """{
    "Code": 605,
    "TimeStamp": 1765036309,
    "Data": {
        "LastIdxSwipe": 30,
        "LastIdxUnkn": 5,
        "LastIdxAdmin": 8,
        "Missing": [
            { "Code": 602, "Index": 31 }
        ]
    }
}""")

    add_table(doc, ["Trường Data", "Backend gửi gì", "Thiết bị làm gì"], [
        ("LastIdxSwipe", "Max Index 602 backend đã lưu (vd 30)", "Gửi lại telemetry 602 có Index > 30 từ rfid_log.csv"),
        ("LastIdxUnkn", "Max Index 601 backend đã lưu", "Gửi lại 601 có Index > LastIdxUnkn"),
        ("LastIdxAdmin", "Max Index 603/604 backend đã lưu", "Gửi lại admin 603/604 từ rfid_log.csv (dòng reg=99|SAVE|index hoặc 99|DEL|index)"),
        ("Missing[]", "Danh sách {Code, Index} thiếu cụ thể", "Gửi lại đúng từng Index (vd thiếu 31 giữa 30 và 32)"),
    ])

    para(doc, "B) Backend GỬI — legacy (chỉ flush queue, không đọc log):", bold=True)
    json_block(doc, """{"Code":605,"TimeStamp":1765036309}""")

    para(doc, "C) Thiết bị TRẢ — response Direct Method (sau khi sync xong, nếu có Data):", bold=True)
    json_block(doc, """{
    "status": 200,
    "payload": {
        "Code": 605,
        "TimeStamp": 1765036309,
        "IdxSwipe": 35,
        "IdxUnkn": 5,
        "IdxAdmin": 8,
        "LastIdxSwipe": 30,
        "PendingFlushed": 0,
        "ResentFromLog": 5,
        "Message": "Sync done"
    }
}""")

    add_table(doc, ["Trường response", "Ý nghĩa cho Backend"], [
        ("IdxSwipe / IdxUnkn / IdxAdmin", "Counter hiện tại trên thiết bị (max index đã cấp)"),
        ("PendingFlushed", "Số message gửi từ az_pend.bin"),
        ("ResentFromLog", "Số message gửi lại từ rfid_log.csv"),
        ("Message", "Sync done hoặc Flush trigger accepted"),
    ])

    para(doc, "D) Thiết bị GỬI LÊN telemetry (Backend NHẬN qua IoT Hub) — KHÔNG phải response 605:", bold=True)
    json_block(doc, """{
    "Code": 602,
    "Index": 31,
    "TimeStamp": 1765036309,
    "Data": {
        "DeviceName": "RFID_Scanner",
        "UID": "09F8D21E",
        "Name": "Nguyen",
        "ID": "NV01"
    }
}""")
    bullet(doc, "Mỗi Index thiếu = 1 message telemetry riêng (31, 32, 33, 34, 35...).")
    bullet(doc, "Index và TimeStamp GIỮ NGUYÊN lúc quẹt — không tạo Index mới.")
    bullet(doc, "Backend dedupe (DeviceId, Code, Index) trước khi insert DB.")

    doc.add_heading("3.4. Code 606 — OTA", level=2)
    para(doc, "Backend GỬI:", bold=True)
    json_block(doc, """{"Code":606,"TimeStamp":1765036309,"Data":{"Url":"https://.../RFID.bin"}}""")
    para(doc, "Thiết bị TRẢ: 200 OTA accepted | 400 Missing Url")

    doc.add_heading("3.5. Edge cases", level=2)
    bullet(doc, "603/604 thiếu UID → thiết bị không response → Azure timeout 30s.")
    bullet(doc, "605 không Data → chỉ flush az_pend, response ngay Flush trigger accepted.")
    bullet(doc, "605 có Data → response TRỄ (sau flush + replay), timeout Azure cần đủ (30s+).")
    bullet(doc, "Không mở IoT Explorer / MQTT client cùng DeviceId (Azure error 400027).")

    # ── 4. Ví dụ đối chiếu 605 ──
    doc.add_heading("4. Ví dụ đối chiếu — Backend phải làm từng bước", level=1)

    doc.add_heading("4.1. Backend có 30, thiết bị có 35 (thiếu 31–35)", level=2)
    num(doc, "Backend lưu last_idx_swipe = 30 cho DeviceId X.")
    num(doc, "Backend gọi invoke-device-method Control:")
    json_block(doc, """{"Code":605,"TimeStamp":1765036309,"Data":{"LastIdxSwipe":30}}""")
    num(doc, "Đợi Direct Method response: ResentFromLog=5, IdxSwipe=35.")
    num(doc, "Đợi thêm telemetry trên Hub: Code 602 Index 31, 32, 33, 34, 35.")
    num(doc, "Backend cập nhật last_idx_swipe = 35. Dedupe nếu trùng.")

    doc.add_heading("4.2. Backend có 30,32,33,34 — thiếu 31 (lỗ hổng)", level=2)
    num(doc, "Backend gọi 605 với Missing hoặc LastIdxSwipe=34 + Missing:")
    json_block(doc, """{"Code":605,"Data":{"Missing":[{"Code":602,"Index":31}]}}""")
    num(doc, "Thiết bị gửi lại telemetry 602 Index=31 từ rfid_log.csv.")
    num(doc, "Backend insert nếu chưa có (DeviceId, 602, 31).")

    doc.add_heading("4.3. Backend có 36, thiết bị IdxSwipe=35 (backend thừa)", level=2)
    bullet(doc, "Response: IdxSwipe=35. Không có message mới.")
    bullet(doc, "Backend tự kiểm tra duplicate / deviceId sai / dữ liệu test.")

    # ── 5. Luồng tổng ──
    doc.add_heading("5. Sơ đồ luồng Backend", level=1)
    para(doc, "Luồng chấm công hàng ngày:", bold=True)
    num(doc, "Nhân viên quét thẻ → Hub nhận telemetry 601 hoặc 602.")
    num(doc, "Backend parse → lưu DB → cập nhật last_idx_*.")
    para(doc, "Luồng đồng bộ khi thiếu:", bold=True)
    num(doc, "Backend so sánh last_idx_* với IdxSwipe trong response trước / nghi gap.")
    num(doc, "Gọi 605 + LastIdx* (+ Missing nếu lỗ hổng giữa).")
    num(doc, "Nhận response 605 → nhận telemetry bù → dedupe → cập nhật last_idx_*.")
    num(doc, "Vẫn thiếu sau 605 → log cảnh báo (log SD trim 60 ngày hoặc message chưa từng quẹt).")

    # ── 6. Bảng mã ──
    doc.add_heading("6. Bảng mã Code tổng hợp", level=1)
    add_table(doc, ["Code", "Hướng", "Ai gửi", "Ai nhận", "Mô tả"], [
        ("601", "Device→Cloud", "Thiết bị", "Backend", "Quét thẻ lạ"),
        ("602", "Device→Cloud", "Thiết bị", "Backend", "Chấm công"),
        ("603", "Cloud→Device", "Backend", "Thiết bị", "Ghi thẻ SD"),
        ("603", "Device→Cloud", "Thiết bị", "Backend", "Xác nhận lưu"),
        ("604", "Cloud→Device", "Backend", "Thiết bị", "Xóa thẻ SD"),
        ("604", "Device→Cloud", "Thiết bị", "Backend", "Xác nhận xóa"),
        ("605", "Cloud→Device", "Backend", "Thiết bị", "Lệnh sync — KHÔNG phải telemetry"),
        ("606", "Cloud→Device", "Backend", "Thiết bị", "OTA"),
    ])

    # ── 7. Lưu trữ SD ──
    doc.add_heading("7. Dữ liệu trên thiết bị (SD card)", level=1)
    add_table(doc, ["File", "Vai trò", "605 dùng?"], [
        ("/sdcard/rfid_log.csv", "Nhật ký mọi sự kiện (quét + admin) — xem mục 7.1", "Có — replay gap"),
        ("/sdcard/az_pend.bin", "Queue chưa publish Hub", "Có — flush bước 1"),
        ("/sdcard/profiles/{UID}.txt", "Profile thẻ", "Không"),
        ("/sdcard/checkin/YYMMDD.txt", "Check-in trong ngày", "Không"),
    ])
    para(doc, "rfid_log.csv giữ ~60 ngày (auto trim). Message cũ hơn không replay được.")

    doc.add_heading("7.1. Định dạng rfid_log.csv", level=2)
    para(doc, "Quét thẻ (601/602 trên Hub):", bold=True)
    json_block(doc, "2025-05-29T08:00:00|09F8D21E|Nguyen|NV01|1|88")
    add_table(doc, ["Cột", "Ý nghĩa"], [
        ("ts", "Thời gian ISO local"),
        ("uid", "UID hex"),
        ("name, id", "Tên/mã NV (rỗng nếu thẻ lạ)"),
        ("reg", "1=đã ĐK → telemetry 602 | 0=thẻ lạ → 601"),
        ("index", "Index telemetry (idx_swipe hoặc idx_unkn)"),
    ])
    para(doc, "Admin lưu/xóa thẻ (603/604 trên Hub):", bold=True)
    json_block(doc, "2025-05-29T08:02:00|09F8D21E|Nguyen|NV01|99|SAVE|3")
    json_block(doc, "2025-05-29T08:03:00|09F8D21E|Nguyen|NV01|99|DEL|4")
    add_table(doc, ["Cột", "Ý nghĩa"], [
        ("reg=99", "Đánh dấu thao tác admin"),
        ("SAVE / DEL", "603 lưu thẻ | 604 xóa thẻ"),
        ("index", "Index telemetry idx_admin (603/604)"),
    ])
    para(doc, "Portal web hiển thị thêm cột Code (601–604) và Index để vận hành đối chiếu với Backend.")

    # ── 8. Azure CLI ──
    doc.add_heading("8. Azure CLI — Backend test", level=1)

    para(doc, "Đăng ký thẻ (603):", bold=True)
    json_block(doc, """az iot hub invoke-device-method \\
  --hub-name {IoTHubName} --device-id {DeviceId} \\
  --method-name Control \\
  --method-payload '{"Code":603,"TimeStamp":1765036309,"Data":{"UID":"09F8D21E","Name":"Nguyen","ID":"NV01"}}'""")

    para(doc, "Xóa thẻ (604):", bold=True)
    json_block(doc, """az iot hub invoke-device-method \\
  --hub-name {IoTHubName} --device-id {DeviceId} \\
  --method-name Control \\
  --method-payload '{"Code":604,"TimeStamp":1765036309,"Data":{"UID":"09F8D21E"}}'""")

    para(doc, "Đồng bộ có đối chiếu (605):", bold=True)
    json_block(doc, """az iot hub invoke-device-method \\
  --hub-name {IoTHubName} --device-id {DeviceId} \\
  --method-name Control \\
  --method-payload '{"Code":605,"TimeStamp":1765036309,"Data":{"LastIdxSwipe":30,"LastIdxUnkn":0,"LastIdxAdmin":0}}'""")

    para(doc, "Bù Index 31 thiếu (605 + Missing):", bold=True)
    json_block(doc, """az iot hub invoke-device-method \\
  --hub-name {IoTHubName} --device-id {DeviceId} \\
  --method-name Control \\
  --method-payload '{"Code":605,"Data":{"Missing":[{"Code":602,"Index":31}]}}'""")

    # ── 9. Checklist ──
    doc.add_heading("9. Checklist triển khai Backend", level=1)
    bullets = [
        "Route telemetry devices/{deviceId}/messages/events/ → worker/DB.",
        "Parse: Code, Index, TimeStamp, Data.UID, Data.Name, Data.ID.",
        "Lưu DB unique key: (DeviceId, Code, Index) — ON CONFLICT ignore hoặc upsert.",
        "Lưu riêng last_idx_swipe (602), last_idx_unkn (601), last_idx_admin (603/604).",
        "Code 602 → ghi chấm công với TimeStamp.",
        "Code 601 → queue/cảnh báo thẻ lạ.",
        "Đăng ký thẻ: invoke 603, đợi response Updated UID + telemetry 603 (Index).",
        "Xóa thẻ: invoke 604, đợi response + telemetry 604 (Index).",
        "Phát hiện gap Index → invoke 605 với LastIdx* và/hoặc Missing[].",
        "Sau 605: đọc response ResentFromLog + nhận telemetry bù trên Hub.",
        "Dedupe mọi message trùng Index khi flush/replay.",
        "KHÔNG connect MQTT bằng DeviceId thiết bị (Explorer/backend service).",
        "Job định kỳ: sync 605 cho thiết bị offline lâu hoặc sau reconnect.",
    ]
    for b in bullets:
        bullet(doc, b)

    doc.add_heading("9.1. Bảng Backend GỬI vs NHẬN", level=2)
    add_table(doc, ["Hành động Backend", "Gửi (Direct Method)", "Nhận từ thiết bị"], [
        ("Chấm công", "— (passive)", "Telemetry 601/602 trên Hub"),
        ("Đăng ký thẻ", "603 + UID,Name,ID", "DM response 200 + telemetry 603 (Index) trên Hub"),
        ("Xóa thẻ", "604 + UID", "DM response 200 + telemetry 604 (Index) trên Hub"),
        ("Sync thiếu", "605 + LastIdx* / Missing", "DM response 605 (ResentFromLog) + telemetry bù trên Hub"),
        ("OTA", "606 + Url", "DM response 200"),
    ])

    return doc


def main():
    doc = build_document()
    for target in [OUTPUT, ROOT / "Maychamcong1.docx"]:
        if target.exists():
            shutil.copy2(target, target.with_suffix(".docx.bak"))
        try:
            doc.save(target)
            print(f"OK: {target}")
        except PermissionError:
            alt = ROOT / (target.stem + "_updated.docx")
            doc.save(alt)
            print(f"LOCKED {target} -> {alt}")


if __name__ == "__main__":
    main()
